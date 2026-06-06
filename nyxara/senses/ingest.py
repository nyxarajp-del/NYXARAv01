"""NYXARA · senses/ingest.py — documents → structured text + RAG-ready chunks (👁).

Knowledge arrives as files, not tidy strings. This sense turns a document of almost any
common shape — plain text, Markdown, CSV/TSV, JSON, source code, HTML, and (when the
optional libraries are present) PDF / DOCX / XLSX — into a single, clean
:class:`Document`: normalised text, structural metadata (columns, headings, language,
line/word counts), and a list of overlapping :class:`Chunk` s sized for retrieval.

Two design commitments:

* **Never crash on a missing library.** Binary formats use ``pdfplumber`` / ``python-docx``
  / ``openpyxl`` *if installed*; if not, the document still ingests — with an honest note
  that the extractor was unavailable — rather than raising. Every text-shaped format works
  on the bare standard library.
* **Boundary-aware chunking.** :class:`Chunker` splits on sentence boundaries (via
  :mod:`senses.nlp`) into character-budgeted chunks with a configurable overlap, and
  records the *exact* source offsets of each chunk, so retrieved passages can be cited
  back to the original.

:meth:`Ingestor.ingest_file` and :meth:`Ingestor.ingest_text` are the entry points; the
result feeds memory and the RAG layer.

Reuses :mod:`senses.nlp` (sentence splitting) and :mod:`senses.web` (HTML). Heavy parsers
are optional. Pure standard library otherwise.
"""

from __future__ import annotations

import csv
import io
import json
import os
import re
from dataclasses import dataclass, field
from enum import Enum
from typing import Any, Dict, List, Optional, Tuple

from nyxara.senses.nlp import sentences

__all__ = [
    "DocumentType",
    "Chunk",
    "Document",
    "Chunker",
    "Ingestor",
]


# --------------------------------------------------------------------------- #
# Types
# --------------------------------------------------------------------------- #
class DocumentType(str, Enum):
    TEXT = "text"
    MARKDOWN = "markdown"
    CSV = "csv"
    TSV = "tsv"
    JSON = "json"
    CODE = "code"
    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    UNKNOWN = "unknown"


_EXT_MAP: Dict[str, DocumentType] = {
    ".txt": DocumentType.TEXT, ".text": DocumentType.TEXT, ".log": DocumentType.TEXT,
    ".md": DocumentType.MARKDOWN, ".markdown": DocumentType.MARKDOWN,
    ".csv": DocumentType.CSV, ".tsv": DocumentType.TSV,
    ".json": DocumentType.JSON, ".jsonl": DocumentType.JSON,
    ".html": DocumentType.HTML, ".htm": DocumentType.HTML,
    ".pdf": DocumentType.PDF, ".docx": DocumentType.DOCX, ".xlsx": DocumentType.XLSX,
    ".py": DocumentType.CODE, ".js": DocumentType.CODE, ".ts": DocumentType.CODE,
    ".java": DocumentType.CODE, ".c": DocumentType.CODE, ".cpp": DocumentType.CODE,
    ".go": DocumentType.CODE, ".rs": DocumentType.CODE, ".rb": DocumentType.CODE,
    ".sh": DocumentType.CODE, ".sql": DocumentType.CODE, ".yaml": DocumentType.CODE,
    ".yml": DocumentType.CODE, ".toml": DocumentType.CODE,
}

_BINARY = {DocumentType.PDF, DocumentType.DOCX, DocumentType.XLSX}


# --------------------------------------------------------------------------- #
# Chunk & document
# --------------------------------------------------------------------------- #
@dataclass
class Chunk:
    index: int
    text: str
    start: int
    end: int
    source: str = ""
    section: str = ""

    @property
    def approx_tokens(self) -> int:
        return max(1, len(self.text) // 4)

    def to_dict(self) -> Dict[str, Any]:
        return {"index": self.index, "text": self.text, "start": self.start,
                "end": self.end, "source": self.source, "section": self.section,
                "approx_tokens": self.approx_tokens}


@dataclass
class Document:
    source: str
    doc_type: DocumentType
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    chunks: List[Chunk] = field(default_factory=list)
    note: str = ""                       # honest note, e.g. a missing optional parser

    @property
    def word_count(self) -> int:
        return len(self.text.split())

    @property
    def char_count(self) -> int:
        return len(self.text)

    @property
    def ok(self) -> bool:
        return bool(self.text) or not self.note

    def to_dict(self) -> Dict[str, Any]:
        return {"source": self.source, "doc_type": self.doc_type.value,
                "word_count": self.word_count, "char_count": self.char_count,
                "n_chunks": len(self.chunks), "metadata": self.metadata,
                "note": self.note}


# --------------------------------------------------------------------------- #
# Chunker (boundary-aware, overlapping, offset-exact)
# --------------------------------------------------------------------------- #
class Chunker:
    """Splits text into overlapping, sentence-aligned chunks with exact source offsets."""

    def __init__(self, *, target_chars: int = 1000, overlap_chars: int = 150) -> None:
        if target_chars <= 0:
            raise ValueError("target_chars must be positive")
        self.target_chars = target_chars
        self.overlap_chars = max(0, min(overlap_chars, target_chars - 1))

    def _sentence_spans(self, text: str) -> List[Tuple[str, int, int]]:
        spans: List[Tuple[str, int, int]] = []
        cursor = 0
        for s in (sentences(text) or ([text] if text.strip() else [])):
            i = text.find(s, cursor)
            if i < 0:
                i = cursor
            spans.append((s, i, i + len(s)))
            cursor = i + len(s)
        return spans

    def chunk(self, text: str, *, source: str = "", section: str = "") -> List[Chunk]:
        text = text or ""
        spans = self._sentence_spans(text)
        if not spans:
            return []
        chunks: List[Chunk] = []
        cur: List[Tuple[str, int, int]] = []
        cur_len = 0
        idx = 0

        def emit() -> None:
            nonlocal idx
            start = cur[0][1]
            end = cur[-1][2]
            chunks.append(Chunk(index=idx, text=text[start:end], start=start, end=end,
                                source=source, section=section))
            idx += 1

        for span in spans:
            s_len = span[2] - span[1]
            if cur and cur_len + s_len > self.target_chars:
                emit()
                # carry trailing sentences as overlap
                keep: List[Tuple[str, int, int]] = []
                klen = 0
                for ss in reversed(cur):
                    sl = ss[2] - ss[1]
                    if klen + sl > self.overlap_chars:
                        break
                    keep.insert(0, ss)
                    klen += sl
                cur, cur_len = keep, klen
            cur.append(span)
            cur_len += s_len
        if cur:
            emit()
        return chunks


# --------------------------------------------------------------------------- #
# Ingestor
# --------------------------------------------------------------------------- #
class Ingestor:
    """Detects a document's type and extracts clean text + structure + chunks."""

    def __init__(self, chunker: Optional[Chunker] = None) -> None:
        self.chunker = chunker or Chunker()

    # ---- type detection ---- #
    @staticmethod
    def detect_type(source: str = "", content: str = "") -> DocumentType:
        ext = os.path.splitext(source)[1].lower()
        if ext in _EXT_MAP:
            return _EXT_MAP[ext]
        sniff = content.lstrip()[:1]
        if sniff in ("{", "["):
            try:
                json.loads(content)
                return DocumentType.JSON
            except (json.JSONDecodeError, ValueError):
                pass
        if re.search(r"(?m)^#{1,6}\s+\S", content):
            return DocumentType.MARKDOWN
        if re.search(r"(?m)^<!doctype html|<html", content, re.I):
            return DocumentType.HTML
        first = content.splitlines()[0] if content else ""
        if "," in first and content.count(",") >= content.count("\n"):
            return DocumentType.CSV
        return DocumentType.TEXT

    # ---- text-shaped extractors ---- #
    def _extract_markdown(self, text: str) -> Tuple[str, Dict[str, Any]]:
        headings = [m.group(2).strip()
                    for m in re.finditer(r"(?m)^(#{1,6})\s+(.+)$", text)]
        clean = re.sub(r"(?m)^#{1,6}\s+", "", text)
        clean = re.sub(r"[*_`>#]+", "", clean)
        clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", clean)  # [text](url) -> text
        return clean.strip(), {"headings": headings, "n_headings": len(headings)}

    def _extract_delimited(self, text: str, delim: str) -> Tuple[str, Dict[str, Any]]:
        reader = list(csv.reader(io.StringIO(text), delimiter=delim))
        if not reader:
            return "", {"rows": 0, "columns": []}
        header = reader[0]
        rows = reader[1:]
        lines = [" | ".join(header)]
        for r in rows:
            lines.append(" | ".join(r))
        meta = {"columns": header, "n_columns": len(header), "rows": len(rows)}
        return "\n".join(lines), meta

    def _extract_json(self, text: str) -> Tuple[str, Dict[str, Any]]:
        try:
            data = json.loads(text)
        except (json.JSONDecodeError, ValueError) as exc:
            return text, {"json_error": str(exc)}
        pretty = json.dumps(data, indent=2, ensure_ascii=False, default=str)
        meta: Dict[str, Any] = {"json_type": type(data).__name__}
        if isinstance(data, dict):
            meta["keys"] = list(data.keys())
        elif isinstance(data, list):
            meta["length"] = len(data)
        return pretty, meta

    def _extract_html(self, text: str) -> Tuple[str, Dict[str, Any]]:
        from nyxara.senses.web import HTMLExtractor  # local import (optional dependency chain)
        ex = HTMLExtractor()
        ex.feed(text)
        return ex.text, {"title": ex.title.strip(), "headings": ex.headings,
                         "n_links": len(ex.links)}

    # ---- binary extractors (optional libs, graceful) ---- #
    def _extract_pdf(self, path: str) -> Tuple[str, Dict[str, Any], str]:
        try:
            import pdfplumber  # type: ignore
        except ImportError:
            return "", {}, "pdfplumber not installed; PDF text not extracted"
        try:
            pages: List[str] = []
            with pdfplumber.open(path) as pdf:  # pragma: no cover - exercised only with lib
                for pg in pdf.pages:
                    pages.append(pg.extract_text() or "")
            return "\n\n".join(pages), {"n_pages": len(pages)}, ""
        except Exception as exc:  # noqa: BLE001
            return "", {}, f"PDF extraction failed: {exc}"

    def _extract_docx(self, path: str) -> Tuple[str, Dict[str, Any], str]:
        try:
            import docx  # type: ignore
        except ImportError:
            return "", {}, "python-docx not installed; DOCX text not extracted"
        try:  # pragma: no cover - exercised only with lib
            d = docx.Document(path)
            paras = [p.text for p in d.paragraphs]
            return "\n".join(paras), {"n_paragraphs": len(paras)}, ""
        except Exception as exc:  # noqa: BLE001
            return "", {}, f"DOCX extraction failed: {exc}"

    def _extract_xlsx(self, path: str) -> Tuple[str, Dict[str, Any], str]:
        try:
            import openpyxl  # type: ignore
        except ImportError:
            return "", {}, "openpyxl not installed; XLSX text not extracted"
        try:  # pragma: no cover - exercised only with lib
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            lines: List[str] = []
            for ws in wb.worksheets:
                for row in ws.iter_rows(values_only=True):
                    lines.append(" | ".join("" if c is None else str(c) for c in row))
            return "\n".join(lines), {"sheets": wb.sheetnames}, ""
        except Exception as exc:  # noqa: BLE001
            return "", {}, f"XLSX extraction failed: {exc}"

    # ---- main entry points ---- #
    def ingest_text(self, text: str, *, source: str = "<text>",
                    doc_type: Optional[DocumentType] = None) -> Document:
        dt = doc_type or self.detect_type(source, text)
        note = ""
        if dt is DocumentType.MARKDOWN:
            clean, meta = self._extract_markdown(text)
        elif dt is DocumentType.CSV:
            clean, meta = self._extract_delimited(text, ",")
        elif dt is DocumentType.TSV:
            clean, meta = self._extract_delimited(text, "\t")
        elif dt is DocumentType.JSON:
            clean, meta = self._extract_json(text)
        elif dt is DocumentType.HTML:
            clean, meta = self._extract_html(text)
        else:  # TEXT, CODE, UNKNOWN
            clean, meta = text.strip(), {}
            if dt is DocumentType.CODE:
                meta = {"lines": text.count("\n") + 1}
        meta.setdefault("lines", text.count("\n") + 1)
        doc = Document(source=source, doc_type=dt, text=clean, metadata=meta, note=note)
        doc.chunks = self.chunker.chunk(clean, source=source)
        return doc

    def ingest_file(self, path: str) -> Document:
        dt = self.detect_type(path)
        if dt in _BINARY:
            if dt is DocumentType.PDF:
                text, meta, note = self._extract_pdf(path)
            elif dt is DocumentType.DOCX:
                text, meta, note = self._extract_docx(path)
            else:
                text, meta, note = self._extract_xlsx(path)
            doc = Document(source=path, doc_type=dt, text=text, metadata=meta, note=note)
            doc.chunks = self.chunker.chunk(text, source=path)
            return doc
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()
        return self.ingest_text(content, source=path, doc_type=dt)


# --------------------------------------------------------------------------- #
# Self-test / demo
# --------------------------------------------------------------------------- #
if __name__ == "__main__":  # pragma: no cover
    import tempfile

    print("=" * 70)
    print("NYXARA document-ingestion self-test")
    print("=" * 70)

    ing = Ingestor(chunker=Chunker(target_chars=120, overlap_chars=30))

    # markdown
    md = "# Title\n\nNYXARA serves the Master. She protects his network.\n\n## Details\nMore text here."
    d = ing.ingest_text(md, source="notes.md")
    print(f"\nmarkdown            : type={d.doc_type.value} headings={d.metadata['headings']}")
    assert d.doc_type is DocumentType.MARKDOWN
    assert d.metadata["headings"] == ["Title", "Details"]
    assert "#" not in d.text

    # csv
    csv_text = "name,role\nJP,Master\nNYXARA,Guardian"
    d = ing.ingest_text(csv_text, source="people.csv")
    print(f"csv                 : columns={d.metadata['columns']} rows={d.metadata['rows']}")
    assert d.metadata["columns"] == ["name", "role"] and d.metadata["rows"] == 2

    # json
    d = ing.ingest_text('{"owner": "JP", "loyal": true}', source="cfg.json")
    print(f"json                : keys={d.metadata['keys']}")
    assert d.doc_type is DocumentType.JSON and d.metadata["keys"] == ["owner", "loyal"]

    # code
    d = ing.ingest_text("def f():\n    return 1\n", source="x.py")
    assert d.doc_type is DocumentType.CODE and d.metadata["lines"] >= 2
    print(f"code                : lines={d.metadata['lines']}")

    # html
    d = ing.ingest_text("<html><h1>Hi</h1><p>Body text.</p></html>", source="p.html")
    assert d.doc_type is DocumentType.HTML and "Body text" in d.text
    print(f"html                : title={d.metadata.get('title')!r} text={d.text!r}")

    # chunking with overlap + exact offsets
    long = ("Sentence one is here. Sentence two follows. Sentence three arrives. "
            "Sentence four appears. Sentence five concludes the passage.")
    d = ing.ingest_text(long, source="doc.txt")
    print(f"\nchunking            : {len(d.chunks)} chunks "
          f"(sizes={[len(c.text) for c in d.chunks]})")
    assert len(d.chunks) >= 2
    # offsets are exact: each chunk's text equals the source slice
    for c in d.chunks:
        assert long[c.start:c.end] == c.text
    # consecutive chunks overlap (share some source range)
    assert d.chunks[1].start < d.chunks[0].end
    print("offset fidelity     : every chunk slice matches the source ✓")
    print("overlap             : consecutive chunks share context ✓")

    # type detection without an extension
    assert Ingestor.detect_type(content='{"a":1}') is DocumentType.JSON
    assert Ingestor.detect_type(content="# Heading\ntext") is DocumentType.MARKDOWN
    print("\nsniff detection     : json/markdown detected from content ✓")

    # graceful binary fallback (no crash if the optional parser is missing)
    tmp = os.path.join(tempfile.gettempdir(), "nyx_fake.pdf")
    with open(tmp, "wb") as f:
        f.write(b"%PDF-1.4 not-a-real-pdf")
    d = ing.ingest_file(tmp)
    print(f"\npdf fallback        : note={d.note!r} (no crash)")
    assert d.doc_type is DocumentType.PDF  # either extracted or noted, never raised
    os.remove(tmp)

    print("\nALL SELF-TESTS PASSED ✓")
